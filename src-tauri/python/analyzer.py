import datetime
import os
import pandas as pd
import matplotlib
# 设置matplotlib为非交互式后端，避免GUI线程问题
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

# 设置中文显示
plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号

class BorrowingAnalyzer:
    def __init__(self):
        """
        初始化借阅分析器
        """
        self.data = None
        self.member_data = None
        self.results = None
    
    def load_excel_files(self, borrow_paths: list[str], member_path: str, progress_callback=None):
        """
        加载Excel数据文件
        :param borrow_paths: 借阅数据文件路径列表
        :param member_path: 读者清单文件路径
        :param progress_callback: 进度回调函数 (total, current, message)
        """
        try:
            # 读取读者清单
            self.member_data = pd.read_excel(member_path)
            
            # 存储所有文件的数据，支持单独分析
            self.file_data_list = []
            self.file_names = []
            
            total_files = len(borrow_paths)
            
            for i, path in enumerate(borrow_paths):
                if progress_callback:
                    progress_callback(total_files, i, f"正在读取文件 {os.path.basename(path)}")
                
                df = pd.read_excel(path)
                self.file_data_list.append(df)
                
                # 提取文件名（不含扩展名）
                file_name = os.path.splitext(os.path.basename(path))[0]
                self.file_names.append(file_name)
                
            # 同时保存合并后的数据（用于整体分析）
            self.data = pd.concat(self.file_data_list, ignore_index=True)
            
            if progress_callback:
                progress_callback(total_files, total_files, "数据加载完成")
                
            return True
            
        except Exception as e:
            raise Exception(f"加载Excel文件失败: {str(e)}")

    def analyze_single_file(self, file_index: int, progress_callback=None):
        """
        分析单个文件的借阅数据
        :param file_index: 文件索引
        :param progress_callback: 进度回调函数 (total, current, message)
        :return: 分析结果对象
        """
        if not hasattr(self, 'file_data_list') or file_index >= len(self.file_data_list):
            raise ValueError("无效的文件索引或未加载数据")
        
        # 临时设置当前分析的数据为指定文件
        original_data = self.data
        self.data = self.file_data_list[file_index]
        
        try:
            # 调用原有的分析方法
            result = self.analyze_with_progress(progress_callback)
            return result
        finally:
            # 恢复原始数据
            self.data = original_data

    def get_file_count(self) -> int:
        """获取加载的文件数量"""
        return len(self.file_data_list) if hasattr(self, 'file_data_list') else 0

    def get_file_name(self, file_index: int) -> str:
        """获取指定索引的文件名"""
        if hasattr(self, 'file_names') and file_index < len(self.file_names):
            return self.file_names[file_index]
        return f"文件{file_index + 1}"

    def analyze_with_progress(self, progress_callback=None):
        """
        分析借阅数据并提供进度回调
        :param progress_callback: 进度回调函数 (total, current, message)
        """
        if self.data is None or self.member_data is None:
            raise ValueError("请先加载数据文件")

        try:
            if progress_callback:
                progress_callback(100, 10, "开始分析数据")

            # 计算借阅时长
            self.data['借阅时长'] = pd.to_datetime(self.data['还书日期']) - pd.to_datetime(self.data['借书日期'])
            self.data['借阅时长'] = self.data['借阅时长'].dt.days
            
            if progress_callback:
                progress_callback(100, 20, "计算借阅频次")

            # 借阅频次
            borrow_frequency = self.data.shape[0]

            if progress_callback:
                progress_callback(100, 30, "计算会员时长")

            # 获取会员时长
            open_date = None
            member_name = self.data['姓名'].iloc[0]
            for i in range(len(self.member_data)):
                if self.member_data['姓名'].iloc[i] == member_name:
                    open_date = self.member_data['办卡日期'].iloc[i]
                    break
            
            if open_date is None:
                total_duration = "未知。请检查读者清单是否包含该读者。"
            else:
                total_duration = datetime.datetime.now() - pd.to_datetime(open_date)

            if progress_callback:
                progress_callback(100, 40, "计算统计指标")

            # 总阅读时长和定价总价
            total_reading_duration = self.data['借阅时长'].sum()
            total_price = self.data['定价'].sum()

            if progress_callback:
                progress_callback(100, 50, "分析月度数据")

            # 月度数据分析
            self.data['借书月份'] = pd.to_datetime(self.data['借书日期']).dt.month
            monthly_duration = self.data.groupby('借书月份')['借阅时长'].sum()
            monthly_duration = monthly_duration.sort_values(ascending=False)
            
            # v2功能：按年-月格式统计月度借阅
            self.data['借书年月'] = pd.to_datetime(self.data['借书日期']).dt.strftime('%Y-%m')
            monthly_borrow = self.data.groupby('借书年月')['书名'].count()
            monthly_borrow = monthly_borrow.sort_values(ascending=False)

            if progress_callback:
                progress_callback(100, 60, "分析图书类别")

            # 类别分析
            self.data['类别'] = self.data['类别'].apply(self.categorycode2name)
            book_duration = self.data.groupby('书名')['借阅时长'].sum()
            book_duration_ratio = book_duration / total_reading_duration * 100
            sorted_duration = book_duration.sort_values(ascending=False)

            # v2新增：计算每本书被借阅的次数（用于文字版排名）
            book_borrow_counts = self.data['书名'].value_counts().sort_values(ascending=False)

            if progress_callback:
                progress_callback(100, 70, "分析最长时间和最频繁统计")

            # 原有：找出阅读时长最长的作者和类别
            longest_author = self.data.groupby('作者')['借阅时长'].sum().idxmax()
            longest_category = self.data.groupby('类别')['借阅时长'].sum().idxmax()
            
            # v2新增：找出借阅次数最多的作者和类别
            most_frequent_author = self.data['作者'].value_counts().idxmax()
            most_frequent_category = self.data['类别'].value_counts().idxmax()

            if progress_callback:
                progress_callback(100, 80, "分析借阅高峰")

            # 借阅高峰期和最多借阅的书
            # v2更新：借阅高峰期改为年-月格式
            borrow_peak_yearmonth = self.data.groupby('借书年月')['书名'].count().sort_values(ascending=False).idxmax()
            # 保留原有的月份高峰期（用于图表）
            borrow_peak = self.data.groupby('借书月份')['书名'].count().sort_values(ascending=False).idxmax()
            most_borrowed_book = self.data['书名'].value_counts().idxmax()
            most_borrowed_book_count = self.data['书名'].value_counts().max()

            if progress_callback:
                progress_callback(100, 90, "统计全勤月")

            # 全勤月统计
            self.data['借书日期'] = pd.to_datetime(self.data['借书日期'])
            self.data['ISO周'] = self.data['借书日期'].apply(
                lambda x: f"{x.isocalendar()[0]}-{x.isocalendar()[1]:02d}")
            self.data['年月'] = self.data['借书日期'].dt.strftime('%Y-%m')
            monthly_weeks = self.data.groupby('年月')['ISO周'].nunique()
            
            # 计算每个月的实际周数
            full_attendance = []
            partial_attendance = []
            missing_weeks_info = {}
            # 获取所有借书记录的周
            all_recorded_weeks = set(self.data['ISO周'].unique())
            for month in monthly_weeks.index:
                year, month_num = map(int, month.split('-'))
                actual_weeks, all_weeks = self.calculate_actual_weeks(year, month_num)
                
                # 获取该月实际借阅的周
                recorded_weeks = self.data[self.data['年月'] == month]['ISO周'].unique()
                
                # 找出缺失的周
                recorded_set = set(recorded_weeks)
                all_set = set(all_weeks)
                
                # 对于跨月周的特殊处理
                # 如果一个周在当月没有记录，但在该周对应的其他月份有记录，则视为有记录
                # 例如，2025-14既是第三月的周也是第四月的周
                # 如果第四月有记录，即使第三月没有记录，2025-14也应被视为第三月有记录
                corrected_recorded_set = set(recorded_set)  # 复制原始记录
                for week in all_set - recorded_set:
                    # 如果该周在其他月份有记录
                    if week in all_recorded_weeks:
                        corrected_recorded_set.add(week)
                
                # 计算修正后的缺失周
                missing_weeks = all_set - corrected_recorded_set
                
                if len(corrected_recorded_set) == actual_weeks:
                    full_attendance.append(month)
                else:
                    partial_attendance.append(month)
                    # 存储缺失周信息
                    missing_weeks_info[month] = {
                        'recorded_weeks': sorted(recorded_weeks),
                        'corrected_recorded_weeks': sorted(corrected_recorded_set),
                        'missing_weeks': sorted(missing_weeks),
                        'total_weeks': sorted(all_weeks)
                    }
            
            total_full_attendance = len(full_attendance)

            # 创建结果对象（兼容Rust PyO3接口）
            class AnalysisResults:
                def __init__(self):
                    self.borrow_frequency = borrow_frequency
                    self.total_duration = str(total_duration)
                    self.total_reading_duration = total_reading_duration
                    self.total_price = total_price
                    self.monthly_borrow = dict(monthly_borrow)
                    self.longest_author = longest_author
                    self.longest_category = longest_category
                    # v2新增字段
                    self.most_frequent_author = most_frequent_author
                    self.most_frequent_category = most_frequent_category
                    self.book_borrow_counts = dict(book_borrow_counts.head(10))
                    self.borrow_peak_yearmonth = borrow_peak_yearmonth
                    # 保留原有字段
                    self.borrow_peak = borrow_peak
                    self.most_borrowed_book = most_borrowed_book
                    self.most_borrowed_book_count = most_borrowed_book_count
                    self.full_attendance = full_attendance
                    self.total_full_attendance = total_full_attendance
            
            # 保存分析结果为字典（用于内部使用）
            self.results = {
                'borrow_frequency': borrow_frequency,
                'total_duration': total_duration,
                'total_reading_duration': total_reading_duration,
                'total_price': total_price,
                'monthly_duration': monthly_duration,
                'monthly_borrow': monthly_borrow,
                'book_duration': book_duration,
                'book_duration_ratio': book_duration_ratio,
                'sorted_duration': sorted_duration,
                # v2新增字段
                'book_borrow_counts': book_borrow_counts,
                'most_frequent_author': most_frequent_author,
                'most_frequent_category': most_frequent_category,
                'borrow_peak_yearmonth': borrow_peak_yearmonth,
                # v2新增：缺勤月文档生成辅助字段
                'partial_attendance': partial_attendance,
                'missing_weeks_info': missing_weeks_info,
                # 保留原有字段
                'longest_author': longest_author,
                'longest_category': longest_category,
                'borrow_peak': borrow_peak,
                'most_borrowed_book': most_borrowed_book,
                'most_borrowed_book_count': most_borrowed_book_count,
                'full_attendance': full_attendance,
                'total_full_attendance': total_full_attendance
            }

            results_obj = AnalysisResults()

            if progress_callback:
                progress_callback(100, 100, "分析完成")

            return results_obj

        except Exception as e:
            raise Exception(f"分析数据失败: {str(e)}")

    def calculate_actual_weeks(self, year: int, month: int) -> tuple[int, list[str]]:
        """
        计算一个月中实际包含的ISO周数
        :param year: 年份
        :param month: 月份
        :return: ISO周数量
        """
        # 创建该月的第一天
        first_day = datetime.date(year, month, 1)
        
        # 计算该月的最后一天
        if month == 12:
            last_day = datetime.date(year + 1, 1, 1) - datetime.timedelta(days=1)
        else:
            last_day = datetime.date(year, month + 1, 1) - datetime.timedelta(days=1)
        
        # 生成该月的所有日期
        dates = [first_day + datetime.timedelta(days=i) 
                 for i in range((last_day - first_day).days + 1)]
        
        # 获取每个日期的ISO周并去重
        iso_weeks = set()
        for d in dates:
            iso_year, iso_week, _ = d.isocalendar()
            iso_weeks.add(f"{iso_year}-{iso_week:02d}")
        
        return len(iso_weeks), sorted(iso_weeks)  # 返回周数和所有周列表

    def categorycode2name(self, code: str) -> str:
        """
        将类别代码转换为类别名称
        :param code: 类别代码
        :return: 类别名称
        """
        category_dict = {
            'BLG': '双语读物',
            'CCM': '中文漫画',
            'CFI': '中文小说',
            'CHP': '中文低幼',
            'CNF': '中文科普',
            'CPB': '中文绘本',
            'CPU': '中文立体读物',
            'CPY': '拼音读物',
            'CRF': '中文家长用书',
            'CYX': '中文音像(书)',
            'ECM': '英文漫画',
            'EER': '英文分级读物',
            'EFI': '英文小说',
            'EHP': '英文低幼',
            'ENF': '英文科普',
            'EPB': '英文绘本',
            'EPU': '英文立体读物',
            'ERF': '英文家长用书',
            'EYX': '英文音像(书)'
        }
        return category_dict.get(code, code)

    def generate_charts(self, output_dir: str, progress_callback=None) -> dict:
        """
        生成所有图表并保存到指定目录
        :param output_dir: 输出目录路径
        :param progress_callback: 进度回调函数 (total, current, message)
        :return: 包含所有图表文件路径的字典
        """
        if self.results is None:
            raise ValueError("请先执行分析")

        try:
            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)
            chart_paths = {}
            
            if progress_callback:
                progress_callback(100, 0, "开始生成图表")

            # 1. 阅读时长条形图
            book_duration = self.results['sorted_duration'].head(10)
            book_duration = book_duration[::-1]
            plt.figure(figsize=(10, 6))
            book_duration.plot(kind='barh')
            plt.title('每本书的阅读时长')
            plt.xlabel('时长（天）')
            plt.ylabel('书名')
            plt.tight_layout()
            plt.savefig(f"{output_dir}/book_duration.png", transparent=True, dpi=150)
            plt.close()
            chart_paths['duration'] = f"{output_dir}/book_duration.png"

            if progress_callback:
                progress_callback(100, 33, "已生成阅读时长图表")

            # 2. 借阅频次折线图
            borrow_per_month = self.data.groupby('借书月份')['书名'].count()
            plt.figure(figsize=(10, 6))
            borrow_per_month.plot(kind='line', marker='o')
            plt.title('每月借阅频次')
            plt.xlabel('月份')
            plt.ylabel('频次')
            plt.tight_layout()
            plt.savefig(f"{output_dir}/borrow_per_month.png", transparent=True, dpi=150)
            plt.close()
            chart_paths['monthly'] = f"{output_dir}/borrow_per_month.png"

            if progress_callback:
                progress_callback(100, 66, "已生成借阅频次图表")

            # 3. 类别阅读册数饼图
            category_count = self.data.groupby('类别')['书名'].count()
            plt.figure(figsize=(8, 8))
            category_count.plot(kind='pie', autopct='%1.1f%%')
            plt.title('类别阅读册数比例')
            plt.ylabel('')
            plt.tight_layout()
            plt.savefig(f"{output_dir}/category_ratio.png", transparent=True, dpi=150)
            plt.close()
            chart_paths['category'] = f"{output_dir}/category_ratio.png"

            if progress_callback:
                progress_callback(100, 100, "图表生成完成")

            return chart_paths

        except Exception as e:
            raise Exception(f"生成图表失败: {str(e)}")

    def export_report(self, output_path: str, progress_callback=None):
        """
        导出分析报告为Word文档
        :param output_path: 输出文件路径
        :param progress_callback: 进度回调函数 (total, current, message)
        """
        if self.results is None:
            raise ValueError("请先执行分析")

        try:
            from docx import Document
            import docx.shared
            doc = Document()

            if progress_callback:
                progress_callback(100, 0, "开始生成报告")

            # 添加标题
            doc.add_heading('读书报告分析结果', level=1)

            # 添加基础统计信息（与原型脚本格式一致）
            if progress_callback:
                progress_callback(100, 25, "添加基础统计信息")

            # v2更新：报告格式与main_v2.py保持一致
            basic_info = (
                f"借阅频次: {self.results['borrow_frequency']}\n"
                f"总时长: {self.results['total_duration']}\n"
                f"总价值: {self.results['total_price']} 元\n"
                f"借阅次数最多的作者: {self.results['most_frequent_author']}\n"
                f"借阅次数最多的类别: {self.results['most_frequent_category']}\n"
                f"阅读时长最长的作者: {self.results['longest_author']}\n"
                f"阅读时长最长的类别: {self.results['longest_category']}\n"
                f"借阅高峰期: {self.results['borrow_peak_yearmonth']}\n"
                f"借阅次数最多的书: {self.results['most_borrowed_book']}，共借阅 {self.results['most_borrowed_book_count']} 次\n"
                f"全勤月总数: {self.results['total_full_attendance']}"
            )
            doc.add_paragraph(basic_info)

            # 添加每月借书数量排序（与原型脚本格式一致）
            if progress_callback:
                progress_callback(100, 40, "添加月度统计")

            doc.add_heading('每月借书数量排序', level=2)
            monthly_borrow_text = '\n'.join([f"{month}: {count}"
                                         for month, count in self.results['monthly_borrow'].items()])
            doc.add_paragraph(monthly_borrow_text)

            # v2更新：添加书籍借阅次数排序（前十）
            if progress_callback:
                progress_callback(100, 55, "添加借阅次数排序")

            doc.add_heading('书籍借阅次数排序（前十）', level=2)
            borrow_counts_text = '\n'.join([f"{book}: {count}次"
                                           for book, count in self.results['book_borrow_counts'].head(10).items()])
            doc.add_paragraph(borrow_counts_text)
            
            # 保留借书时长排序（用于图表）
            doc.add_heading('借书时长排序（前十）', level=2)
            sorted_duration_text = '\n'.join([f"{book}: {duration}天"
                                           for book, duration in self.results['sorted_duration'].head(10).items()])
            doc.add_paragraph(sorted_duration_text)

            # 添加全勤月列表
            if progress_callback:
                progress_callback(100, 65, "添加全勤月列表")

            doc.add_heading('全勤月列表', level=2)
            doc.add_paragraph(', '.join(self.results['full_attendance']))

            # 添加缺勤月列表
            if progress_callback:
                progress_callback(100, 70, "添加缺勤月列表")
            
            doc.add_heading('缺勤月列表', level=2)
            if self.results['partial_attendance']:
                partical_attendance_info = ""
                for month in self.results['partial_attendance']:
                    missing_info = self.results['missing_weeks_info'].get(month, {})
                    partical_attendance_info += f"{month}:\n"
                    partical_attendance_info += f"  实际借阅周: {', '.join(missing_info.get('recorded_weeks', []))}\n"
                    partical_attendance_info += f"  修正后借阅周: {', '.join(missing_info.get('corrected_recorded_weeks', []))}\n"
                    partical_attendance_info += f"  缺失周: {', '.join(missing_info.get('missing_weeks', []))}\n"
                    partical_attendance_info += f"  总周数: {', '.join(missing_info.get('total_weeks', []))}\n"
                doc.add_paragraph(partical_attendance_info.strip())
            else:
                doc.add_paragraph("无缺勤月")

            # 添加图表
            if progress_callback:
                progress_callback(100, 75, "生成并添加图表")

            # 确保输出目录存在
            chart_dir = os.path.dirname(output_path)
            if not chart_dir:
                chart_dir = "."
            
            chart_paths = self.generate_charts(chart_dir)

            # 每本书的阅读时长
            doc.add_heading('每本书的阅读时长', level=2)
            doc.add_picture(chart_paths['duration'],
                          width=docx.shared.Inches(6),
                          height=docx.shared.Inches(4))
            
            # 每月借阅频次
            doc.add_heading('每月借阅频次', level=2)
            doc.add_picture(chart_paths['monthly'],
                          width=docx.shared.Inches(6),
                          height=docx.shared.Inches(4))
            
            # 类别阅读时长比例
            doc.add_heading('类别阅读时长比例', level=2)
            doc.add_picture(chart_paths['category'],
                          width=docx.shared.Inches(6),
                          height=docx.shared.Inches(4))

            # 保存文档
            if progress_callback:
                progress_callback(100, 90, "保存报告")

            doc.save(output_path)

            if progress_callback:
                progress_callback(100, 100, "报告生成完成")

        except Exception as e:
            raise Exception(f"导出报告失败: {str(e)}")

    def export_reports_for_all_files(self, base_output_dir: str, progress_callback=None):
        """
        为所有加载的文件分别生成报告（模拟原型脚本的循环处理逻辑）
        :param base_output_dir: 基础输出目录
        :param progress_callback: 进度回调函数 (total, current, message)
        :return: 生成的报告文件路径列表
        """
        if not hasattr(self, 'file_data_list'):
            raise ValueError("请先加载数据文件")

        try:
            report_paths = []
            total_files = len(self.file_data_list)
            
            for i in range(total_files):
                if progress_callback:
                    progress_callback(total_files, i, f"正在为文件 {self.get_file_name(i)} 生成报告")
                
                # 为每个文件创建独立的输出目录
                file_name = self.get_file_name(i)
                file_output_dir = os.path.join(base_output_dir, file_name)
                os.makedirs(file_output_dir, exist_ok=True)
                
                # 分析单个文件
                self.analyze_single_file(i)
                
                # 生成报告
                report_path = os.path.join(file_output_dir, f"{file_name}_analysis.docx")
                self.export_report(report_path)
                
                report_paths.append(report_path)
            
            if progress_callback:
                progress_callback(total_files, total_files, "所有报告生成完成")
            
            return report_paths

        except Exception as e:
            raise Exception(f"批量生成报告失败: {str(e)}")