import datetime
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.animation as animation
import numpy as np


# 设置中文显示
plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号

# 读取数据
class BorrowingAnalyzer:
    def __init__(self, data, member_data):
        """
        初始化借阅分析器
        :param data: 借阅数据的DataFrame
        """
        self.data = data
        self.member_data = member_data
        self.results = None

    def analyze(self):
        """
        分析借阅数据
        """
        # 计算借阅时长
        self.data['借阅时长'] = pd.to_datetime(self.data['还书日期']) - pd.to_datetime(self.data['借书日期'])
        self.data['借阅时长'] = self.data['借阅时长'].dt.days

        # 借阅频次
        borrow_frequency = self.data.shape[0]

        # 总时长
        # 改为会员时长
        open_date = None
        member_name = self.data['姓名'].iloc[0]
        print(member_name)
        for i in range(len(self.member_data)):
            print(self.member_data['姓名'].iloc[i])
            if self.member_data['姓名'].iloc[i] == member_name:
                open_date = self.member_data['办卡日期'].iloc[i]
                break
        
        if open_date == None:
            total_duration = "未知。请检查读者清单是否包含该读者。"
            print("未找到该读者的办卡日期，请检查读者清单是否包含该读者。")
        else:
            total_duration = datetime.datetime.now() - pd.to_datetime(open_date)

        # 总阅读时长
        total_reading_duration = self.data['借阅时长'].sum()

        # 定价总价
        total_price = self.data['定价'].sum()

        # 计算每个月的借书时长
        self.data['借书月份'] = pd.to_datetime(self.data['借书日期']).dt.month
        monthly_duration = self.data.groupby('借书月份')['借阅时长'].sum()
        # 月份排序
        monthly_duration = monthly_duration.sort_values(ascending=False)

        # 每月借书数量排序
        # 修改月份统计，添加年份
        self.data['借书年月'] = pd.to_datetime(self.data['借书日期']).dt.strftime('%Y-%m')
        monthly_borrow = self.data.groupby('借书年月')['书名'].count()
        monthly_borrow = monthly_borrow.sort_values(ascending=False)

        # 小标识换为中文
        self.data['类别'] = self.data['类别'].apply(self.categorycode2name)

        # 每种书的阅读时长
        book_duration = self.data.groupby('书名')['借阅时长'].sum()

        # 阅读时长比例
        book_duration_ratio = book_duration / total_reading_duration * 100

        # 借书时长排序
        sorted_duration = book_duration.sort_values(ascending=False) # 降序排列

        # === 计算每本书被借阅的次数（用于文字版排名）===
        book_borrow_counts = self.data['书名'].value_counts().sort_values(ascending=False)

        # 阅读时长最长的作者和类别
        longest_author = self.data.groupby('作者')['借阅时长'].sum().idxmax()
        longest_category = self.data.groupby('类别')['借阅时长'].sum().idxmax()
        # 改为借阅次数最多的作者和类别
        most_frequent_author = self.data['作者'].value_counts().idxmax()
        most_frequent_category = self.data['类别'].value_counts().idxmax()

        # 借阅高峰期（按月份统计）
        self.data['借书月份'] = pd.to_datetime(self.data['借书日期']).dt.month
        borrow_peak = self.data.groupby('借书月份')['书名'].count().sort_values(ascending=False).idxmax()

        # 搜索借阅次数最多的书
        most_borrowed_book = self.data['书名'].value_counts().idxmax()
        most_borrowed_book_count = self.data['书名'].value_counts().max()

        # ===== 全勤月统计功能 =====
        # 确保借书日期是datetime类型
        self.data['借书日期'] = pd.to_datetime(self.data['借书日期'])
        
        # 添加ISO周列 (格式: YYYY-WW)
        self.data['ISO周'] = self.data['借书日期'].apply(
            lambda x: f"{x.isocalendar()[0]}-{x.isocalendar()[1]:02d}")
        
        # 添加年月列 (格式: YYYY-MM)
        self.data['年月'] = self.data['借书日期'].dt.strftime('%Y-%m')
        
        # 按年月分组，统计每个月的不同周数
        monthly_weeks = self.data.groupby('年月')['ISO周'].nunique()
        
        # 计算每个月的实际周数
        full_attendance = []
        partial_attendance = []
        missing_weeks_info = {}
        
        # 创建一个字典来存储每个周出现在哪些月份
        week_to_months = {}
        for month in monthly_weeks.index:
            year, month_num = map(int, month.split('-'))
            actual_weeks, all_weeks = self.calculate_actual_weeks(year, month_num)
            
            # 记录每个周对应的月份
            for week in all_weeks:
                if week not in week_to_months:
                    week_to_months[week] = set()
                week_to_months[week].add(month)
        
        # 获取所有借书记录的周
        all_recorded_weeks = set(self.data['ISO周'].unique())
        
        for month in monthly_weeks.index:
            year, month_num = map(int, month.split('-'))
            actual_weeks, all_weeks = self.calculate_actual_weeks(year, month_num)
            
            # 获取该月实际借阅的周
            recorded_weeks = self.data[self.data['年月'] == month]['ISO周'].unique()
            
            # 找出缺失的周
            recorded_set = set(recorded_weeks)
            all_set = set(all_weeks)
            
            # 对于跨月周的特殊处理
            # 如果一个周在当月没有记录，但在该周对应的其他月份有记录，则视为有记录
            # 例如，2025-14既是第三月的周也是第四月的周
            # 如果第四月有记录，即使第三月没有记录，2025-14也应被视为第三月有记录
            corrected_recorded_set = set(recorded_set)  # 复制原始记录
            for week in all_set - recorded_set:
                # 如果该周在其他月份有记录
                if week in all_recorded_weeks:
                    corrected_recorded_set.add(week)
            
            # 计算修正后的缺失周
            missing_weeks = all_set - corrected_recorded_set
            
            if len(corrected_recorded_set) == actual_weeks:
                full_attendance.append(month)
            else:
                partial_attendance.append(month)
                # 存储缺失周信息
                missing_weeks_info[month] = {
                    'recorded_weeks': sorted(recorded_weeks),
                    'corrected_recorded_weeks': sorted(corrected_recorded_set),
                    'missing_weeks': sorted(missing_weeks),
                    'total_weeks': sorted(all_weeks)
                }
        
        total_full_attendance = len(full_attendance)
        # ===== 全勤月统计结束 =====

        self.results = {
            'borrow_frequency': borrow_frequency,
            'total_duration': total_duration,
            'total_price': total_price,
            'monthly_duration': monthly_duration,
            'monthly_borrow': monthly_borrow,
            'book_duration': book_duration,
            'book_duration_ratio': book_duration_ratio,
            'sorted_duration': sorted_duration,
            'book_borrow_counts': book_borrow_counts,
            'longest_author': longest_author,
            'longest_category': longest_category,
            'most_frequent_author': most_frequent_author,
            'most_frequent_category': most_frequent_category,
            'borrow_peak': borrow_peak,
            'most_borrowed_book': most_borrowed_book,
            'most_borrowed_book_count': most_borrowed_book_count,
            'full_attendance': full_attendance,
            'partial_attendance': partial_attendance,
            'missing_weeks_info': missing_weeks_info,
            'total_full_attendance': total_full_attendance
        }

        # 格式化输出
        output = (
            f"借阅频次: {borrow_frequency}\n"
            f"总时长: {total_duration} 天\n"
            f"总价值: {total_price} 元\n"
            # f"每月借书时长排序（每个都换行）: \n{monthly_duration}\n"
            f"每月借书数量排序（每个都换行）: \n{monthly_borrow}\n"
            # f"每种书的阅读时长: {book_duration.to_dict()}\n"
            # f"阅读时长比例: {book_duration_ratio.to_dict()}\n"
            # 格式化输出前十个
            f"书籍借阅次数排序（前十）: \n{book_borrow_counts.head(10)}\n"
            f"借书时长排序（前十，每个都换行）: \n{sorted_duration.head(10)}\n"
            f"阅读次数最多的作者: {most_frequent_author}\n"
            f"阅读次数最多的类别: {most_frequent_category}\n"
            f"阅读时长最长的作者: {longest_author}\n"
            f"阅读时长最长的类别: {longest_category}\n"
            f"借阅高峰期: {borrow_peak} 月\n"
            f"借阅次数最多的书: {most_borrowed_book}，共借阅 {most_borrowed_book_count} 次\n"
            f"全勤月列表: {full_attendance}\n"
            f"全勤月总数: {total_full_attendance}"
        )
        # 新增：添加缺勤月详细信息
        if partial_attendance:
            output += "\n\n缺勤月信息:"
            for month in partial_attendance:
                info = missing_weeks_info[month]
                output += (
                    f"\n{month}: 应包含周数: {len(info['total_weeks'])}, "
                    f"实际借阅周数: {len(info['corrected_recorded_weeks'])}, "
                    f"缺失周数: {len(info['missing_weeks'])}"
                )
                output += f"\n  实际借阅周: {', '.join(info['corrected_recorded_weeks'])}"
                output += f"\n  缺失周: {', '.join(info['missing_weeks'])}"
                output += f"\n  所有周: {', '.join(info['total_weeks'])}\n"
        else:
            output += "\n\n没有缺勤月"
        return output

    def calculate_actual_weeks(self, year, month):
        """计算一个月中实际包含的ISO周数及所有周列表"""
        # 创建该月的第一天
        first_day = datetime.date(year, month, 1)
        
        # 计算该月的最后一天
        if month == 12:
            last_day = datetime.date(year + 1, 1, 1) - datetime.timedelta(days=1)
        else:
            last_day = datetime.date(year, month + 1, 1) - datetime.timedelta(days=1)
        
        # 生成该月的所有日期
        dates = [first_day + datetime.timedelta(days=i) 
                 for i in range((last_day - first_day).days + 1)]
        
        # 获取每个日期的ISO周并去重
        iso_weeks = set()
        for d in dates:
            iso_year, iso_week, _ = d.isocalendar()
            iso_weeks.add(f"{iso_year}-{iso_week:02d}")
        
        return len(iso_weeks), sorted(iso_weeks)  # 返回周数和所有周列表

    def categorycode2name(self, code):
        """
        将类别代码转换为类别名称
        """
        # 类别代码和名称的映射关系
        # BLG.双语读物
        # CCM.中文漫画
        # CFI.中文小说
        # CHP.中文低幼CNF.中文科普CPB.中文绘本
        # CPU.中文立体读物
        # CPY.拼音读物
        # CRF.中文家长用书CYX.中文音像(书)ECM.英文漫画
        # EER.英文分级读物
        # EFI.英文小说
        # EHP.英文低幼
        # ENF.英文科普
        # EPB.英文绘本
        # EPU.英文立体读物
        # ERF.英文家长用书
        # EYX.英文音像(书)
        category_dict = {
            'BLG': '双语读物',
            'CCM': '中文漫画',
            'CFI': '中文小说',
            'CHP': '中文低幼',
            'CNF': '中文科普',
            'CPB': '中文绘本',
            'CPU': '中文立体读物',
            'CPY': '拼音读物',
            'CRF': '中文家长用书',
            'CYX': '中文音像(书)',
            'ECM': '英文漫画',
            'EER': '英文分级读物',
            'EFI': '英文小说',
            'EHP': '英文低幼',
            'ENF': '英文科普',
            'EPB': '英文绘本',
            'EPU': '英文立体读物',
            'ERF': '英文家长用书',
            'EYX': '英文音像(书)'
        }
        return category_dict.get(code, code)

    def plot_duration(self, show=True):
        """
        绘制每本书（前十）的阅读时长条形图
        """
        if self.results is None:
            raise ValueError("请先执行分析方法 'analyze'")

        book_duration = self.results['sorted_duration'].head(10)
        book_duration = book_duration[::-1]  # Reverse the order of the DataFrame
        plt.figure(figsize=(10, 6))
        book_duration.plot(kind='barh')  # Change 'kind' to 'barh' for horizontal bar chart
        plt.title('每本书的阅读时长')
        plt.xlabel('时长（天）')  # Swap x-axis and y-axis labels
        plt.ylabel('书名')  # Swap x-axis and y-axis labels
        if show:
            plt.show()
        return plt
    
    def plot_duration_gif(self, show=True):
        """
        绘制每本书（前十）的阅读时长条形图GIF
        """
        if self.results is None:
            raise ValueError("请先执行分析方法 'analyze'")

        book_duration = self.results['sorted_duration'].head(10)
        book_duration = book_duration[::-1]  # Reverse the order of the DataFrame

        fig, ax = plt.subplots(figsize=(10, 6))
        bars = ax.barh(book_duration.index, [0] * len(book_duration), color='skyblue')
        ax.set_title('每本书的阅读时长')
        ax.set_xlabel('时长（天）')
        ax.set_ylabel('书名')
        ax.set_xlim(0, book_duration.max() + 1)

        def init():
            for bar in bars:
                bar.set_width(0)
            return bars

        def update(frame):
            for i, bar in enumerate(bars):
                if i <= frame:
                    bar.set_width(book_duration.values[i])
            return bars

        ani = animation.FuncAnimation(fig, update, frames=len(bars), init_func=init, blit=True, repeat=False)

        if show:
            plt.show()

        return ani

    def plot_books_per_month(self, show=True):
        """
        绘制借阅频次折线图
        """
        if self.results is None:
            raise ValueError("请先执行分析方法 'analyze'")
        
        borrow_per_month = self.data.groupby('借书月份')['书名'].count()
        plt.figure(figsize=(10, 6))
        borrow_per_month.plot(kind='line', marker='o')
        plt.title('每月借阅频次')
        plt.xlabel('月份')
        plt.ylabel('频次')
        if show:
            plt.show()
        return plt
        
    def plot_books_per_month_gif(self, show=True):
        """
        绘制借阅频次折线图并保存为GIF
        """
        if self.results is None:
            raise ValueError("请先执行分析方法 'analyze'")
        
        borrow_per_month = self.data.groupby('借书月份')['书名'].count()
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.set_title('每月借阅频次')
        ax.set_xlabel('月份')
        ax.set_ylabel('频次')
        ax.set_xlim(1, 12)
        ax.set_ylim(0, borrow_per_month.max() + 1)
        
        line, = ax.plot([], [], marker='o')

        def init():
            line.set_data([], [])
            return line,

        def update(frame):
            x = borrow_per_month.index[:frame + 1]
            y = borrow_per_month.values[:frame + 1]
            line.set_data(x, y)
            return line,

        ani = animation.FuncAnimation(fig, update, frames=len(borrow_per_month), init_func=init, blit=True, repeat=False)

        if show:
            plt.show()
        
        return ani

    def plot_category(self, show=True):
        """
        绘制类别阅读册数的饼图
        """
        if self.results is None:
            raise ValueError("请先执行分析方法 'analyze'")

        categories = self.data['类别'].unique()
        # category_duration = self.data.groupby('类别')['借阅时长'].sum()
        category_count = self.data.groupby('类别')['书名'].count()
        plt.figure(figsize=(8, 8))
        # category_duration.plot(kind='pie', autopct='%1.1f%%')
        category_count.plot(kind='pie', autopct='%1.1f%%')
        plt.title('类别阅读册数比例')
        plt.ylabel('')
        if show:
            plt.show()
        return plt
    
    def plot_category_gif(self, show=True):
        """
        绘制类别阅读册数的饼图gif
        """
        if self.results is None:
            raise ValueError("请先执行分析方法 'analyze'")

        category_count = self.data.groupby('类别')['书名'].count()
        categories = category_count.index
        sizes = category_count.values

        fig, ax = plt.subplots(figsize=(8, 8))
        wedges, texts, autotexts = ax.pie(sizes, labels=categories, autopct='%1.1f%%', startangle=90)
        ax.set_title('类别阅读册数比例')

        def init():
            for wedge in wedges:
                wedge.set_visible(False)
            return wedges

        def update(frame):
            for i, wedge in enumerate(wedges):
                if i <= frame:
                    wedge.set_visible(True)
            return wedges

        ani = animation.FuncAnimation(fig, update, frames=len(wedges), init_func=init, blit=True, repeat=False)


        if show:
            plt.show()
        
        return ani



# 读取数据
# GUI中选择文件
from tkinter import filedialog
from tkinter import Tk

# Create the Tkinter root
root = Tk()
root.withdraw()  # Hide the main window

# Ask the user to select SEVERAL files
# file_path = filedialog.askopenfilename(title="选择借阅数据文件", filetypes=[("Excel files", ["*.xlsx", "*.xls"])])
file_paths = filedialog.askopenfilenames(title="选择多个借阅数据文件", filetypes=[("Excel files", ["*.xlsx", "*.xls"])])
file_path2 = filedialog.askopenfilename(title="选择读者清单", filetypes=[("Excel files", ["*.xlsx", "*.xls"])])
df2 = pd.read_excel(file_path2)
# 序号 账号 卡号 姓名 性别 类别 状态 办卡日期 截止日期 操作日期
# 1 1001 1001 张三 男 会员 有效 2021-01-01 2022-01-01 2021-01-01

for file_path in file_paths:
    # print(file_path)

    # Read the Excel file
    df = pd.read_excel(file_path)
    # Convert '借书日期' and '还书日期' to datetime
    # df['借书日期'] = pd.to_datetime(df['借书日期'])
    # df['还书日期'] = pd.to_datetime(df['还书日期'])

    # Calculate the borrowing duration and add it to the DataFrame
    # df['借阅时长'] = (df['还书日期'] - df['借书日期']).dt.days

    # Drop the '日期差' column as it is no longer needed
    # df.drop(columns=['日期差'], inplace=True)

    # Display the updated DataFrame
    df.head()

    #-------------------------

    # 使用示例数据进行分析和绘图
    analyzer = BorrowingAnalyzer(df, df2)
    formatted_results = analyzer.analyze()
    print(formatted_results)
    # analyzer.plot_duration()
    # # analyzer.plot_ratio()
    # analyzer.plot_books_per_month()
    # analyzer.plot_category()

    # 保存分析结果到docx
    from docx import Document
    import docx
    import os

    file_name = file_path.split('/')[-1]      # .split('.')[0]
    save_folder = f"../../output/{file_name}/"
    if not os.path.exists(save_folder):
        os.makedirs(save_folder)
    doc = Document()
    doc.add_heading('读书报告分析结果', level=1)
    doc.add_paragraph(formatted_results)

    doc.add_heading('每本书的阅读时长', level=2)
    analyzer.plot_duration(show=False)
    plt.savefig(save_folder + 'book_duration.png', transparent=True)
    plt.close()

    doc.add_picture(save_folder + 'book_duration.png', width=docx.shared.Inches(6), height=docx.shared.Inches(4))

    doc.add_heading('每月借阅频次', level=2)
    analyzer.plot_books_per_month(show=False)
    plt.savefig(save_folder + 'borrow_per_month.png', transparent=True)
    plt.close()

    doc.add_picture(save_folder + 'borrow_per_month.png', width=docx.shared.Inches(6), height=docx.shared.Inches(4))

    doc.add_heading('类别阅读时长比例', level=2)
    analyzer.plot_category(show=False)
    plt.savefig(save_folder + 'category_duration.png', transparent=True)
    plt.close()

    doc.add_picture(save_folder + 'category_duration.png', width=docx.shared.Inches(6), height=docx.shared.Inches(4))

    doc.save(save_folder + f'{file_name}_analysis.docx')

    # 保存GIF
    # book_duration_gif = analyzer.plot_duration_gif(show=False)
    # book_duration_gif.save(save_folder + 'book_duration.gif', writer='pillow', fps=2)

    # books_per_month_gif = analyzer.plot_books_per_month_gif(show=False)
    # books_per_month_gif.save(save_folder + 'borrow_per_month.gif', writer='pillow', fps=2)

    # category_gif = analyzer.plot_category_gif(show=False)
    # category_gif.save(save_folder + 'category_duration.gif', writer='pillow', fps=2)

    print("分析结果已保存到：", save_folder + f'{file_name}_analysis.docx')